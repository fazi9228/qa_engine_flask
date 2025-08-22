import re
import io
import os
import hashlib
import datetime
from typing import List, Dict, Any, Optional
import time
import tempfile
import glob
import unicodedata

# Try to import file handling libraries
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

class EnhancedChatProcessor:
    """Enhanced processor for extracting and parsing chat transcripts from various file formats"""

    def extract_chats_from_file(self, uploaded_file):
        """
        Extract chats from uploaded file based on file type
        
        Args:
        uploaded_file: The uploaded file object (BytesIO with .name attribute or file-like object)
        
        Returns:
        List of chat dictionaries
        """
        try:
            # Handle both file path strings and file objects
            if hasattr(uploaded_file, 'name'):
                filename = uploaded_file.name
            elif hasattr(uploaded_file, 'filename'):
                filename = uploaded_file.filename
            else:
                raise ValueError("File object must have 'name' or 'filename' attribute")
                
            file_extension = filename.split('.')[-1].lower()
            print(f"Processing file: {filename}, Type: {file_extension}")
            
            # Always reset file pointer to beginning
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
            
            if file_extension == 'txt':
                return self._extract_from_txt(uploaded_file)
            elif file_extension == 'csv':
                if pd is None:
                    raise ImportError("pandas library is required for CSV processing. Install with: pip install pandas")
                return self._extract_from_csv(uploaded_file)
            elif file_extension == 'pdf':
                if PyPDF2 is None:
                    raise ImportError("PyPDF2 library is required for PDF processing. Install with: pip install PyPDF2")
                return self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                if docx is None:
                    raise ImportError("python-docx library is required for DOCX processing. Install with: pip install python-docx")
                return self._extract_from_docx(uploaded_file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            import traceback
            print(f"Error in extract_chats_from_file: {str(e)}")
            print(traceback.format_exc())
            return []

    def _extract_from_txt(self, uploaded_file):
        """Extract chats from a text file"""
        try:
            # Reset file pointer to beginning
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
            
            # Handle different file object types
            if hasattr(uploaded_file, 'read'):
                # For BytesIO objects or file-like objects
                content_bytes = uploaded_file.read()
                if isinstance(content_bytes, bytes):
                    content = content_bytes.decode('utf-8')
                else:
                    content = content_bytes
            else:
                raise ValueError("File object must have a 'read' method")
                
            return self._split_text_into_chats(content)
        except Exception as e:
            print(f"Error extracting from TXT: {str(e)}")
            return []

    def _extract_from_csv(self, uploaded_file):
        """Extract chats from a CSV file"""
        try:
            # Reset file pointer to beginning
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
            
            # Read CSV into pandas DataFrame
            # Handle both BytesIO and file-like objects
            if hasattr(uploaded_file, 'getvalue'):
                # For BytesIO objects
                csv_content = uploaded_file.getvalue()
                if isinstance(csv_content, bytes):
                    csv_content = csv_content.decode('utf-8')
                df = pd.read_csv(io.StringIO(csv_content))
            else:
                # For other file-like objects
                df = pd.read_csv(uploaded_file)

            # Check for common columns that might contain chat content
            content_columns = [col for col in df.columns if any(
                term in col.lower() for term in ['chat', 'message', 'content', 'text', 'transcript']
            )]

            if not content_columns:
                # If no obvious content column, use the first column containing text data
                for col in df.columns:
                    if df[col].dtype == 'object':
                        content_columns = [col]
                        break

            # If still no content column found, use all columns
            if not content_columns:
                print("Warning: Couldn't identify chat content columns. Using all columns.")
                content_columns = df.columns.tolist()

            # Join all content into a single string
            content = ""

            if len(content_columns) == 1:
                # Use entire CSV as one chat
                main_col = content_columns[0]
                content = "\n".join(df[main_col].astype(str).tolist())
            else:
                # Try to format as a dialogue
                for _, row in df.iterrows():
                    for col in content_columns:
                        content += f"{col}: {row[col]}\n"
                    content += "\n"

            return self._split_text_into_chats(content)

        except Exception as e:
            print(f"Error extracting from CSV: {str(e)}")
            return []
    
    def _extract_from_pdf(self, uploaded_file):
        """Extract chats from a PDF file"""
        try:
            # Reset file pointer to beginning
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
            
            # Handle different file object types
            if hasattr(uploaded_file, 'getvalue'):
                # For BytesIO objects
                pdf_data = uploaded_file.getvalue()
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
            else:
                # For other file-like objects
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                
            content = ""

            # Extract text from all pages
            for page_num in range(len(pdf_reader.pages)):
                content += pdf_reader.pages[page_num].extract_text() + "\n\n"
                
            return self._split_text_into_chats(content)
            
        except Exception as e:
            print(f"Error extracting from PDF: {str(e)}")
            return []

    def _extract_from_docx(self, uploaded_file):
        """
        Extract chats from a DOCX file, now with support for tables and improved error handling.
    
        """
        try:
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
            
            doc = docx.Document(uploaded_file)
            full_text = []

            # Iterate through paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text = para.text.strip()
                    # Normalize Unicode text to handle mixed encodings
                    text = unicodedata.normalize('NFC', text)
                    full_text.append(text)

            # Iterate through tables to extract all text from cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if the cell has text to avoid errors
                        if cell.text:
                            full_text.append(cell.text.strip())
            
            # Join all extracted text into a single block
            content = "\n".join(full_text)
            
            if not content:
                print("Warning: No text content was extracted from the DOCX file.")
                return []

            return self._split_text_into_chats(content)
            
        except Exception as e:
            import traceback
            print(f"Error extracting from DOCX: {str(e)}")
            print(traceback.format_exc())
            return []

    def _extract_conversation_id_and_type(self, header_text: str) -> tuple[Optional[str], str]:
        """
        Extract conversation ID and determine if it's a Chat or Case.
        FINAL VERSION: Preserves original prefixes like 'MS-' when detected.
    
        """
        # NEW: First, check for the specific 'PREFIX-NUMBER' format (e.g., MS-00148928)
        # The pattern captures the prefix and the number in two separate groups
        prefix_match = re.search(r'\b([A-Z]{2,}-)(\d+)\b', header_text)
        if prefix_match:
            prefix = prefix_match.group(1) # e.g., "MS-"
            number = prefix_match.group(2) # e.g., "00148928"
            # Determine type based on context, default to 'chat' if unclear
            convo_type = "case" if "case" in header_text.lower() else "chat"
            return f"{prefix}{number}", convo_type

        # --- Fallback to original logic for all other formats ---

        # Chat patterns
        chat_patterns = [
            r"Chat\s*#?\s*(\d+)",
            r"Chat\s*:\s*(\d+)",
            r"Chat\s+ID\s*:?\s*(\d+)",
            r"[A-Z]{2,}\s*[:-]?\s*Chat\s*#?\s*(\d+)",
        ]
        
        # Case patterns
        case_patterns = [
            r"Case\s+ID\s+CT(\d+)",
            r"Case\s+(\d+)",
            r"Case\s*#\s*(\d+)",
            r"Case\s*:\s*(\d+)",
        ]
        
        # Check for Chat patterns
        for pattern in chat_patterns:
            match = re.search(pattern, header_text, re.IGNORECASE)
            if match:
                chat_id = match.group(1)
                return f"Chat_{chat_id}", "chat"
        
        # Check for Case patterns
        for pattern in case_patterns:
            match = re.search(pattern, header_text, re.IGNORECASE)
            if match:
                case_id = match.group(1)
                # Special handling for "CT" prefix
                if "CT" in match.group(0).upper():
                    return f"Case_CT{case_id}", "case"
                return f"Case_{case_id}", "case"
        
        # If no match found, generate a hash as a fallback
        hash_obj = hashlib.md5(header_text.encode())
        return f"Unknown_{hash_obj.hexdigest()[:8]}", "unknown"
    
    def _split_text_into_chats(self, text: str) -> List[Dict[str, Any]]:
        """
        Split a text containing multiple chat/case transcripts into individual conversations.
        FINAL VERSION: Flexibly finds all known header formats to reliably split chats.
    
        """
        conversations = []
        
        MIN_LINES = 3
        MIN_CHARS = 50
        
        # This list contains all valid patterns that can start a new conversation.
        conversation_start_patterns = [
            r'\b[A-Z]{2,}-\d+\b', # For MS-00148928
            r"Chat\s*#?\s*\d+",
            r"Chat\s*:\s*\d+",
            r"Chat\s+ID\s*:?\s*\d+",
            r"[A-Z]{2,}\s*[:-]?\s*Chat\s*#?\s*\d+",
            r"Case\s+ID\s+CT\d+",
            r"Case\s+\d+",
            r"Case\s*#\s*\d+",
            r"Case\s*:\s*\d+",
        ]
        
        combined_pattern = '|'.join(f"(?:{p})" for p in conversation_start_patterns)
        
        # Use re.finditer to get the start and end position of each header match.
        # This is more robust than splitting the string.
        matches = list(re.finditer(combined_pattern, text, re.IGNORECASE | re.MULTILINE))
        
        if not matches:
            # Fallback for files with no recognizable headers, treat as one chat.
            if len(text.strip()) > MIN_CHARS:
                return [{
                    'id': 'Conversation_Single',
                    'type': 'unknown',
                    'content': text,
                    'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'processed_content': self._clean_and_process_chat(text)
                }]
            return []

        # Iterate through the matches to slice the text into individual chats.
        for i, match in enumerate(matches):
            start_pos = match.start()
            # The end of the current chat is the start of the next one.
            end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            
            full_conversation_text = text[start_pos:end_pos].strip()
            header = match.group(0).strip()
            
            lines = [line for line in full_conversation_text.split('\n') if line.strip()]
            if len(lines) >= MIN_LINES and len(full_conversation_text) >= MIN_CHARS:
                conversation_id, conversation_type = self._extract_conversation_id_and_type(header)
                
                conversations.append({
                    'id': conversation_id or f"Unknown_{i}",
                    'type': conversation_type,
                    'content': full_conversation_text,
                    'timestamp': self._extract_timestamp(full_conversation_text) or datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'processed_content': self._clean_and_process_chat(full_conversation_text)
                })
                
        return conversations

    def _extract_chat_id(self, chat_text: str) -> Optional[str]:
        """Extract chat/case ID from text - Updated for backward compatibility."""
        conversation_id, _ = self._extract_conversation_id_and_type(chat_text)
        return conversation_id

    def _extract_timestamp(self, chat_text: str) -> Optional[str]:
        """Extract timestamp from chat text."""
        timestamp_patterns = [
            r"Chat Started:\s*([^,\n]+(?:\s+\d{4})?)",
            r"Session Started:\s*([^,\n]+(?:\s+\d{4})?)",
            r"Conversation Started:\s*([^,\n]+(?:\s+\d{4})?)"
        ]

        for pattern in timestamp_patterns:
            match = re.search(pattern, chat_text, re.IGNORECASE)
            if match:
                try:
                    timestamp_str = match.group(1).strip()
                    # Try to parse the timestamp - handle various formats
                    for fmt in [
                        "%A, %B %d, %Y, %H:%M:%S",
                        "%A, %B %d, %Y %H:%M:%S",
                        "%Y-%m-%d %H:%M:%S"
                    ]:
                        try:
                            parsed_time = datetime.datetime.strptime(timestamp_str, fmt)
                            return parsed_time.strftime("%Y-%m-%d %H:%M:%S")
                        except ValueError:
                            continue
                except Exception:
                    pass
        return None

    def _clean_and_process_chat(self, chat_text: str) -> str:
        """
        Clean and format chat content for analysis - UPDATED to preserve category headers
        This version handles inconsistent line breaks AND preserves important metadata.
        """
        # === NEW: Pre-processing step to normalize inconsistent line breaks ===
        original_lines = chat_text.split('\n')
        normalized_lines = []
        i = 0
        while i < len(original_lines):
            line = original_lines[i].strip()
            
            # IMPORTANT: Preserve category headers - don't process them as speaker lines
            if line.startswith('**Chat reason:') or line.startswith('Chat reason:') or 'Chat reason:' in line:
                normalized_lines.append(line)
                i += 1
                continue
                
            # Check if the NEXT line starts with the '•' timestamp marker
            if (i + 1) < len(original_lines) and original_lines[i+1].strip().startswith('•'):
                # This is a split speaker/timestamp. Join them.
                speaker_name = line
                timestamp_part = original_lines[i+1].strip()
                # Combine them into a single, standardized line
                normalized_lines.append(f"{speaker_name} {timestamp_part}")
                i += 2 # Skip the next line since we've already processed it
            else:
                # This is a regular line
                if line:
                    normalized_lines.append(line)
                i += 1
        # =======================================================================

        processed_lines = []
        has_valid_content = False
        
        # UPDATED: Skip patterns that DON'T include category headers
        skip_patterns = [
            r'^\s*\{ChatWindowButton:',
            r'^\s*Agent joined the conversation\.',
            r'^\s*Automated Process\s*•',
            r'^\s*A transfer request was sent',
            r'^.*(left|joined) the conversation\s*•',
            r'^\s*Preview:.*\.(jpg|png|pdf|jpeg)',
            # REMOVED: The general asterisk pattern that was catching category headers
            # r'^\s*[-=*_]{3,}\s*$',  # This was removing **Chat reason: ...**
        ]
        
        # Patterns are now simpler because the input is normalized
        speaker_patterns = {
            'customer': [
                r'^Guest\s*•',
            ],
            'agent': [
                r'^[A-Za-z\s]+\s*•', # Matches any name (e.g., "Jeremy N", "Kang A") followed by '•'
            ]
        }
        
        # For backward compatibility with old formats that use colons
        old_format_speaker_regex = re.compile(r'^(?:Customer|Client|Visitor|User|Guest|Agent|Support|Bot|Pepper):', re.IGNORECASE)

        # Compile patterns for the new format
        customer_regex = re.compile('|'.join(speaker_patterns['customer']), re.IGNORECASE)
        agent_regex = re.compile('|'.join(speaker_patterns['agent']), re.IGNORECASE)
        skip_regex = re.compile('|'.join(skip_patterns), re.IGNORECASE)

        current_speaker = None
        current_message_lines = []

        def commit_message():
            nonlocal has_valid_content
            if current_speaker and current_message_lines:
                message_text = ' '.join(current_message_lines).strip()
                if message_text:
                    # For old formats, the speaker tag is already in the message, so we don't add it again.
                    if not old_format_speaker_regex.match(message_text):
                        processed_lines.append(f"{current_speaker}: {message_text}")
                    else:
                        processed_lines.append(message_text)
                    has_valid_content = True
            current_message_lines.clear()

        # Process the NORMALIZED lines
        for line in normalized_lines:
            # IMPORTANT: Preserve category headers as-is
            if 'Chat reason:' in line or 'Category:' in line or 'Issue type:' in line or 'Topic:' in line:
                processed_lines.append(line)  # Keep category headers intact
                has_valid_content = True
                continue
                
            if skip_regex.match(line):
                continue

            is_new_speaker = False
            if customer_regex.search(line):
                commit_message()
                current_speaker = "Customer"
                is_new_speaker = True
            elif agent_regex.search(line):
                commit_message()
                current_speaker = "Agent"
                is_new_speaker = True
            elif old_format_speaker_regex.match(line):
                # Handle old formats
                commit_message()
                speaker_tag = line.split(':')[0]
                current_speaker = "Agent" if "agent" in speaker_tag.lower() or "support" in speaker_tag.lower() else "Customer"
                current_message_lines.append(line) # Keep the original line with the tag
                is_new_speaker = True

            if not is_new_speaker and current_speaker:
                current_message_lines.append(line)

        commit_message()

        if has_valid_content:
            return '\n'.join(processed_lines)
        
        return ''

def cleanup_session_state():
    """Clean up old session data"""
    max_age = 3600  # 1 hour
    now = time.time()
    # Note: This function needs to be adapted for Flask session management
    pass

class RateLimiter:
    def __init__(self, calls_per_minute=60):
        self.calls_per_minute = calls_per_minute
        self.calls = []
    
    def wait_if_needed(self):
        now = time.time()
        self.calls = [t for t in self.calls if now - t < 60]
        if len(self.calls) >= self.calls_per_minute:
            time.sleep(60 - (now - self.calls[0]))
        self.calls.append(now)

def cleanup_temp_files():
    """Clean up temporary audio files"""
    temp_dir = tempfile.gettempdir()
    pattern = "qa_engine_*.wav"
    for f in glob.glob(os.path.join(temp_dir, pattern)):
        try:
            os.remove(f)
        except OSError:
            pass
